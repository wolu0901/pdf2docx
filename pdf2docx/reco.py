# -*- coding: utf-8 -*-
"""
Created on Mon Oct 30 22:40:59 2017

@author: wolu0
"""
#lang='chi_sim'
from PIL import Image
import pytesseract
import docx
import pdfread
import os

filename='Final MS.pdf'
shortfile=filename.replace('.pdf','')

def reco(targetfile,imagepath):
    text=pytesseract.image_to_string(Image.open(imagepath), lang='eng')
    print(text)
    inputdocx(text,targetfile)
    
  #rwdox/  
#    with open('a.docx','a+',encoding='utf-8') as f:
#        f.write(text)
#        r = f.readlines()
#        print (r)


def inputdocx(text,filename):   
    doc=docx.Document(filename)
    doc.add_paragraph(text)
    doc.save(filename)

def createdocx(filename):
    if not os.path.exists(filename):
        doc_new=docx.Document()
        doc_new.save(filename)
    else:
        print('docx file already exists and I will replace it.')

def createfolder(filename):
    if not os.path.exists(filename):
        os.mkdir(filename)
    else:
        print('folder already exists')
    
def main():
    createfolder(shortfile)
    pdfread.convert_pdf_to_jpg(filename,shortfile)
    targetfile='{}.docx'.format(shortfile)
    createdocx(targetfile)

    for i in range (0, pdfread.convert_pdf_to_jpg(filename,shortfile)-1):
        reco(targetfile,'{0}/{1}-{2}.jpeg'.format(shortfile,shortfile,i))
        
main()
    